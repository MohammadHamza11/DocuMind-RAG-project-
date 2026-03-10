"""Document loading and chunking utilities."""

from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument

from app.config import settings


def extract_text_from_pdf(file_path: Path) -> list[Document]:
    """Extract text from a PDF file, one Document per page."""
    reader = PdfReader(str(file_path))
    documents = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "page": page_num,
                    },
                )
            )
    return documents


def extract_text_from_docx(file_path: Path) -> list[Document]:
    """Extract text from a DOCX file, one Document per paragraph group."""
    doc = DocxDocument(str(file_path))
    full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    if not full_text.strip():
        return []
    return [
        Document(
            page_content=full_text,
            metadata={"source": file_path.name, "page": 1},
        )
    ]


def load_document(file_path: Path) -> list[Document]:
    """Load a document (PDF or DOCX) and return raw Document objects."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def chunk_documents(documents: list[Document]) -> list[Document]:
    """Split documents into smaller chunks for embedding."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(documents)
